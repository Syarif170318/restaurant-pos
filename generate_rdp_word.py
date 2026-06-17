# -*- coding: utf-8 -*-
"""
Generate dokumen Word RDP berformat formal dari file Markdown.
Seluruh bab, tabel, diagram, dan lampiran dikonversi secara lengkap.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).parent
MD_FILE = BASE_DIR / "RDP_Restaurant_Cafe_POS.md"
OUT_FILE = BASE_DIR / "RDP_Restaurant_Cafe_POS.docx"

COLOR_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
COLOR_SECONDARY = RGBColor(0x37, 0x41, 0x51)
COLOR_HEADER_BG = "1E40AF"
COLOR_ALT_ROW = "F1F5F9"
COLOR_CODE_BG = "F8FAFC"
COLOR_ACCENT = RGBColor(0x0F, 0x76, 0x6E)
COLOR_NOTE_BG = "EFF6FF"
COLOR_TIP_BG = "F0FDF4"

# Uraian singkat pembuka per bab — gaya dokumen formal
CHAPTER_NOTES: dict[str, str] = {
    "1. Pendahuluan": (
        "Bab ini memuat latar belakang proyek, tujuan sistem, "
        "serta definisi istilah utama yang digunakan dalam dokumen."
    ),
    "2. Ruang Lingkup Sistem": (
        "Bab ini mendefinisikan ruang lingkup fungsional sistem, "
        "fitur yang berada di luar cakupan (out of scope), "
        "serta asumsi dan ketergantungan operasional."
    ),
    "3. Stakeholder & Peran Pengguna": (
        "Bab ini menguraikan pemangku kepentingan, pembagian peran pengguna, "
        "dan matriks hak akses. Kasir dan waiter mengoperasikan aplikasi mobile; "
        "admin, manager, inventory, dan kitchen mengoperasikan dashboard web."
    ),
    "4. Alur Bisnis (Business Process)": (
        "Bab ini mendeskripsikan lima alur bisnis utama: pemesanan, dapur, "
        "pembayaran, inventori, serta pelaporan dan analitik."
    ),
    "5. Kebutuhan Fungsional": (
        "Bab ini memuat daftar kebutuhan fungsional beserta identifikasi unik (ID) "
        "untuk keperluan pelacakan implementasi dan pengujian."
    ),
    "6. Kebutuhan Non-Fungsional": (
        "Bab ini menetapkan standar performa, ketersediaan, skalabilitas, "
        "keamanan, usability, dan kompatibilitas sistem."
    ),
    "7. Arsitektur Sistem": (
        "Bab ini menguraikan arsitektur teknis sistem. "
        "Versi 1.1 mengadopsi arsitektur dual-client: Dashboard Web dan Mobile App "
        "yang terintegrasi pada satu backend API."
    ),
    "8. Desain Database": (
        "Bab ini mendefinisikan struktur basis data, relasi entitas, "
        "serta spesifikasi tabel untuk transaksi, menu, pembayaran, dan inventori."
    ),
    "9. Desain Modul & API": (
        "Bab ini memuat spesifikasi modul aplikasi dan struktur endpoint API "
        "untuk komunikasi antara klien (web/mobile) dan server."
    ),
    "10. Desain Antarmuka (UI/UX)": (
        "Bab ini memuat daftar halaman aplikasi dan wireframe antarmuka pengguna "
        "sebagai acuan desain visual."
    ),
    "11. Integrasi Eksternal": (
        "Bab ini mendefinisikan integrasi dengan perangkat dan layanan pihak ketiga, "
        "termasuk printer thermal, terminal kartu, gateway QRIS, dan layanan email."
    ),
    "12. Keamanan & Hak Akses": (
        "Bab ini mengatur mekanisme autentikasi, otorisasi berbasis peran, "
        "serta kebijakan audit trail untuk aktivitas kritis."
    ),
    "13. Aturan Bisnis (Business Rules)": (
        "Bab ini menetapkan aturan bisnis yang wajib dipatuhi oleh sistem, "
        "meliputi format transaksi, validasi order, perhitungan tagihan, dan inventori."
    ),
    "14. Strategi Pengujian": (
        "Bab ini mendefinisikan strategi dan skenario pengujian "
        "untuk memastikan kualitas sistem sebelum deployment."
    ),
    "15. Deployment & Operasional": (
        "Bab ini menguraikan lingkungan deployment, infrastruktur minimum, "
        "serta kebutuhan perangkat di outlet operasional."
    ),
    "16. Rencana Pengembangan (Milestone)": (
        "Bab ini memuat rencana pengembangan bertahap, status implementasi per fase, "
        "serta deliverable yang direncanakan."
    ),
    "17. Glosarium": (
        "Bab ini memuat definisi istilah teknis dan bisnis yang digunakan dalam dokumen."
    ),
    "18. Lampiran": (
        "Bab ini memuat lampiran pendukung: diagram status, format transaksi, "
        "rumus perhitungan, checklist review, dan arsitektur dual-client."
    ),
}


def set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_paragraph_shading(paragraph, fill_hex: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("RDP — Restaurant / POS System  |  v1.1")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    hr.font.italic = True

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Halaman ").font.size = Pt(9)

    run = fp.add_run()
    for el_type, text in [("begin", None), (None, "PAGE"), ("separate", None), ("end", None)]:
        if text:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), el_type)
            run._r.append(fld)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_SECONDARY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.2

    for level, size, color, space_before in [
        (1, 18, COLOR_PRIMARY, 18),
        (2, 14, COLOR_PRIMARY, 12),
        (3, 12, COLOR_ACCENT, 8),
        (4, 11, COLOR_SECONDARY, 6),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        add_header_footer(section)


def parse_inline_formatting(paragraph, text: str, base_size=11, base_font="Calibri"):
    pattern = re.compile(r"(\*\*.*?\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            run.font.size = Pt(base_size)
            run.font.name = base_font

        token = match.group()
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.size = Pt(base_size)
            run.font.name = base_font
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xBE, 0x12, 0x3C)
        pos = match.end()

    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(base_size)
        run.font.name = base_font


def add_info_box(doc: Document, title: str, body: str, bg: str = COLOR_NOTE_BG):
    p_title = doc.add_paragraph()
    set_paragraph_shading(p_title, bg)
    p_title.paragraph_format.left_indent = Cm(0.4)
    p_title.paragraph_format.right_indent = Cm(0.4)
    p_title.paragraph_format.space_before = Pt(6)
    r = p_title.add_run(title)
    r.bold = True
    r.font.color.rgb = COLOR_PRIMARY
    r.font.size = Pt(11)

    p_body = doc.add_paragraph()
    set_paragraph_shading(p_body, bg)
    p_body.paragraph_format.left_indent = Cm(0.4)
    p_body.paragraph_format.right_indent = Cm(0.4)
    p_body.paragraph_format.space_after = Pt(8)
    parse_inline_formatting(p_body, body, base_size=10)


def add_cover_page(doc: Document):
    for _ in range(5):
        doc.add_paragraph()

    for text, size, color, bold in [
        ("RDP", 36, COLOR_PRIMARY, True),
        ("Restaurant / POS System", 22, COLOR_SECONDARY, True),
        ("Rancangan Desain Perangkat Lunak", 16, COLOR_ACCENT, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.bold = bold
        run.font.name = "Calibri"
        if not bold and size == 16:
            run.italic = True

    doc.add_paragraph()
    info_data = [
        ("Nama Proyek", "Restaurant / POS System"),
        ("Versi Dokumen", "1.1"),
        ("Tanggal", "8 Juni 2026"),
        ("Status", "Draft v1.1 — Review Stakeholder"),
        ("Referensi", "Workflow Diagram — Restaurant / POS System Workflow"),
        ("Distribusi Utama", "Project Owner, Project Manager, Tim Developer, Tim QA"),
    ]
    table = doc.add_table(rows=len(info_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_data):
        c0, c1 = table.rows[i].cells[0], table.rows[i].cells[1]
        c0.paragraphs[0].add_run(label).bold = True
        c0.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY
        c1.paragraphs[0].add_run(value)
        if i % 2 == 0:
            set_cell_shading(c0, COLOR_ALT_ROW)
            set_cell_shading(c1, COLOR_ALT_ROW)

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = foot.add_run("Dokumen ini bersifat rahasia dan untuk keperluan internal proyek.")
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    doc.add_page_break()


def add_executive_summary(doc: Document):
    doc.add_heading("Ringkasan Eksekutif", level=1)
    add_info_box(
        doc,
        "Tujuan Dokumen",
        "Dokumen Rancangan Desain Perangkat Lunak (RDP) ini merupakan acuan resmi "
        "perancangan sistem Point of Sale (POS) untuk restoran dan kafe. "
        "Dokumen ditujukan bagi pemangku kepentingan bisnis dan tim teknis "
        "sebagai dasar review, validasi kebutuhan, dan persetujuan pengembangan.",
    )

    sections = [
        (
            "Ruang Lingkup Sistem",
            "Sistem POS terintegrasi yang mencakup manajemen pemesanan, Kitchen Display System (KDS), "
            "pemrosesan pembayaran, pengelolaan inventori bahan baku, "
            "serta pelaporan dan analitik penjualan harian.",
        ),
        (
            "Arsitektur Dual-Client",
            "Sistem diimplementasikan dalam dua aplikasi klien yang terintegrasi pada satu backend:\n\n"
            "1. Dashboard Web (pos-app) — dioperasikan oleh Admin, Manager, Inventory, "
            "dan Kitchen Staff melalui peramban web pada PC atau laptop.\n"
            "2. Mobile Application (pos-mobile) — dioperasikan oleh Kasir dan Waiter "
            "pada perangkat Android dan iOS.\n\n"
            "Kedua aplikasi klien menggunakan basis data dan server API yang sama.",
        ),
        (
            "Alur Bisnis Utama",
            "Sistem mengakomodasi lima alur bisnis: Order, Kitchen, Payment, Inventory, "
            "dan Reports & Analytics. Setiap alur terintegrasi secara otomatis "
            "untuk menjaga konsistensi data transaksi dan persediaan.",
        ),
        (
            "Status Implementasi",
            "Fase 1 (MVP) dan Fase 2 (Enhancement) telah diselesaikan. "
            "Arsitektur dual-client (Dashboard Web dan Mobile Application) telah diimplementasikan. "
            "Fase 3 (Scale) — meliputi multi-outlet, program loyalitas, "
            "dan menu digital pelanggan — belum dimulai.",
        ),
        (
            "Ruang Review Stakeholder",
            "Stakeholder diminta melakukan review terhadap alur bisnis (Bab 4), "
            "kebutuhan fungsional (Bab 5), pembagian peran (Bab 3), "
            "serta rencana pengembangan (Bab 16). "
            "Catatan revisi dapat dicantumkan pada Formulir Review Stakeholder "
            "di bagian Kontrol Dokumen.",
        ),
    ]
    for title, body in sections:
        doc.add_heading(title, level=2)
        for para in body.split("\n\n"):
            p = doc.add_paragraph()
            parse_inline_formatting(p, para.strip())

    doc.add_page_break()


def add_reading_guide(doc: Document):
    doc.add_heading("Panduan Pembacaan Dokumen", level=1)

    guide_rows = [
        ["Peran Pembaca", "Bab yang Direkomendasikan"],
        ["Project Owner / Pemilik Bisnis", "Ringkasan Eksekutif, Bab 1, 4, 16"],
        ["Manager Operasional", "Bab 3, 4, 5, 13, 16"],
        ["Kasir / Waiter (End-User)", "Bab 4.1–4.3, Bab 10, Lampiran F"],
        ["Kitchen Staff", "Bab 4.2, Bab 10.3"],
        ["Inventory Staff", "Bab 4.4, Bab 5.6, Bab 10"],
        ["Tim IT / Developer", "Bab 7, 8, 9, 12, 15"],
        ["Tim QA", "Bab 5, 13, 14"],
    ]
    add_table(doc, guide_rows)

    add_info_box(
        doc,
        "Keterangan Prioritas Kebutuhan Fungsional (Bab 5)",
        "Must — kebutuhan wajib pada rilis awal.  |  "
        "Should — kebutuhan yang sangat direkomendasikan.  |  "
        "Could — kebutuhan opsional yang dapat ditunda ke fase berikutnya.  "
        "Setiap kebutuhan memiliki kode ID (contoh: FR-ORD-001) "
        "untuk pelacakan implementasi dan pengujian.",
        bg=COLOR_TIP_BG,
    )

    term_rows = [
        ["Istilah", "Definisi"],
        ["POS", "Point of Sale; sistem kasir dan manajemen transaksi penjualan"],
        ["KDS", "Kitchen Display System; antarmuka digital pesanan di area dapur"],
        ["Dashboard", "Aplikasi web manajemen, pelaporan, dan konfigurasi sistem"],
        ["Mobile App", "Aplikasi native Android/iOS untuk operasional kasir dan waiter"],
        ["Modifier", "Opsi tambahan pada item menu (tingkat pedas, ukuran, dan sejenisnya)"],
        ["Void", "Pembatalan order sebelum pembayaran dilakukan"],
        ["Refund", "Pengembalian dana setelah transaksi pembayaran"],
        ["Shift", "Periode operasional kasir, meliputi pembukaan dan penutupan modal"],
        ["BOM", "Bill of Materials; komposisi bahan baku per item menu"],
        ["QRIS", "Quick Response Code Indonesian Standard; metode pembayaran digital"],
        ["Split Bill", "Pembagian tagihan berdasarkan item atau porsi pembayaran"],
    ]
    doc.add_heading("Daftar Istilah Utama", level=2)
    add_table(doc, term_rows)
    doc.add_page_break()


def add_toc(doc: Document, toc_items: list[str]):
    doc.add_heading("Daftar Isi", level=1)

    main_items = []
    for item in toc_items:
        m = re.match(r"^(\d+)\.\s+(.+)$", item)
        if m:
            num, title = m.groups()
            title = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", title)
            main_items.append((num, title))

    for num, title in main_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{num}. ")
        r1.bold = True
        r1.font.color.rgb = COLOR_PRIMARY
        p.add_run(title)

    doc.add_paragraph()
    doc.add_heading("Lampiran (Bab 18)", level=2)
    for letter, desc in [
        ("A", "Diagram status order"),
        ("B", "Diagram status item di dapur"),
        ("C", "Format nomor transaksi"),
        ("D", "Rumus perhitungan tagihan"),
        ("E", "Checklist review dokumen"),
        ("F", "Arsitektur dual-client (Dashboard Web + Mobile Application)"),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        parse_inline_formatting(p, f"**Lampiran {letter}** — {desc}")

    doc.add_page_break()


def add_document_control(doc: Document):
    doc.add_heading("Kontrol Dokumen", level=1)

    doc.add_heading("Riwayat Revisi", level=2)
    add_table(
        doc,
        [
            ["Versi", "Tanggal", "Penulis", "Keterangan Perubahan"],
            ["1.0", "8 Juni 2026", "Tim Pengembang", "Rilis awal berdasarkan workflow diagram"],
            [
                "1.1",
                "8 Juni 2026",
                "Tim Pengembang",
                "Fase 2, arsitektur dual-client, status implementasi, Lampiran F",
            ],
        ],
    )

    doc.add_heading("Status Implementasi", level=2)
    add_table(
        doc,
        [
            ["Komponen", "Status", "Keterangan"],
            [
                "Fase 1 — MVP",
                "Selesai",
                "Order, KDS, pembayaran tunai/kartu, inventori, laporan harian",
            ],
            [
                "Fase 2 — Enhancement",
                "Selesai",
                "Modifier, split bill, QRIS, void, shift management, stock opname",
            ],
            [
                "Dashboard Web",
                "Selesai",
                "Pelaporan, manajemen menu, inventori, supplier, konfigurasi, KDS",
            ],
            [
                "Mobile Application (Android/iOS)",
                "Selesai",
                "Pemesanan, pembayaran, dan manajemen shift untuk kasir dan waiter",
            ],
            [
                "Fase 3 — Scale",
                "Belum Dimulai",
                "Multi-outlet, program loyalitas, menu digital pelanggan, integrasi delivery",
            ],
        ],
    )

    doc.add_heading("Distribusi Dokumen", level=2)
    add_table(
        doc,
        [
            ["Penerima", "Tujuan Distribusi", "Tingkat Akses"],
            ["Project Owner", "Pengambilan keputusan bisnis dan persetujuan proyek", "Full"],
            ["Manager Operasional", "Validasi kesesuaian alur operasional outlet", "Full"],
            ["Tim Developer", "Implementasi dan pemeliharaan sistem", "Full"],
            ["Tim QA", "Perencanaan dan eksekusi pengujian", "Bab 5, 13, 14"],
            ["Stakeholder Bisnis", "Review kebutuhan dan ruang lingkup", "Ringkasan Eksekutif, Bab 4, 16"],
        ],
    )

    doc.add_heading("Formulir Review Stakeholder", level=2)
    add_info_box(
        doc,
        "Petunjuk Pengisian",
        "Stakeholder mengisi kolom Persetujuan (Ya/Tidak) dan mencantumkan "
        "catatan revisi pada kolom Keterangan. "
        "Formulir diserahkan kepada tim proyek sebelum dimulainya pengembangan Fase 3.",
    )
    review_rows = [
        ["No", "Aspek Review", "Persetujuan (Ya/Tidak)", "Keterangan / Usulan Revisi"],
        [
            "1",
            "Kesesuaian alur bisnis pemesanan, dapur, dan pembayaran dengan operasional outlet",
            "",
            "",
        ],
        [
            "2",
            "Kesesuaian pembagian fungsi Dashboard Web dan Mobile Application",
            "",
            "",
        ],
        [
            "3",
            "Kelengkapan dan kecukupan daftar kebutuhan fungsional (Bab 5)",
            "",
            "",
        ],
        [
            "4",
            "Kesesuaian matriks hak akses per peran pengguna (Bab 3)",
            "",
            "",
        ],
        [
            "5",
            "Kecukupan metode pembayaran yang didukung (tunai, kartu, QRIS)",
            "",
            "",
        ],
        [
            "6",
            "Kesesuaian kebutuhan pelaporan dan analitik (Bab 5.7) dengan kebutuhan bisnis",
            "",
            "",
        ],
        [
            "7",
            "Prioritas fitur Fase 3 yang perlu didahulukan (Bab 16)",
            "",
            "",
        ],
        [
            "8",
            "Kebutuhan tambahan yang belum tercakup dalam dokumen",
            "",
            "",
        ],
    ]
    add_table(doc, review_rows)
    doc.add_page_break()


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s:|\-]+\|?$", line.strip()))


def parse_table_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return

    headers = rows[0]
    data_rows = rows[1:]
    num_cols = len(headers)
    font_size = 8 if num_cols >= 7 else (9 if num_cols >= 5 else 10)

    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, COLOR_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.strip(":"))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(font_size)

    for i, row_data in enumerate(data_rows):
        for j in range(num_cols):
            cell = table.rows[i + 1].cells[j]
            text = row_data[j] if j < len(row_data) else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if num_cols <= 4 else WD_ALIGN_PARAGRAPH.CENTER
            parse_inline_formatting(p, text, base_size=font_size)
            if i % 2 == 1:
                set_cell_shading(cell, COLOR_ALT_ROW)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_code_block(doc: Document, lines: list[str], lang: str = ""):
    if lang == "json":
        label = doc.add_paragraph()
        run = label.add_run("Contoh struktur data JSON:")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    p = doc.add_paragraph()
    set_paragraph_shading(p, COLOR_CODE_BG)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    for idx, line in enumerate(lines):
        if idx > 0:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8 if len(lines) > 15 else 9)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)


def add_blockquote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_paragraph_shading(p, COLOR_NOTE_BG)
    parse_inline_formatting(p, text, base_size=10)


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.4)
    parse_inline_formatting(p, text)


def collect_blockquote(lines: list[str], start: int) -> tuple[str, int]:
    parts: list[str] = []
    i = start
    while i < len(lines):
        s = lines[i].strip()
        if s.startswith(">"):
            content = s[1:].strip()
            if content:
                parts.append(content)
            i += 1
        elif not s:
            if i + 1 < len(lines) and lines[i + 1].strip().startswith(">"):
                i += 1
            else:
                break
        else:
            break
    return "\n".join(parts), i


def is_persona_line(text: str) -> bool:
    return bool(re.match(r"^\*\*Persona \d+", text))


def is_label_line(text: str) -> bool:
    return bool(re.match(r"^\*\*[^*]+\*\*:?\s*$", text))


def count_md_sections(lines: list[str]) -> list[str]:
    sections = []
    for line in lines:
        s = line.strip()
        if s.startswith("## ") and not s.startswith("## Daftar") and not s.startswith("## Rancangan"):
            sections.append(s[3:].strip())
    return sections


def convert_markdown_to_docx():
    lines = MD_FILE.read_text(encoding="utf-8").splitlines()
    expected_sections = count_md_sections(lines)

    doc = Document()
    setup_styles(doc)
    add_cover_page(doc)

    toc_items = []
    in_toc = False
    for line in lines:
        if line.strip() == "## Daftar Isi":
            in_toc = True
            continue
        if in_toc:
            if line.strip().startswith("---"):
                break
            if re.match(r"^\d+\.", line.strip()):
                toc_items.append(line.strip())

    add_toc(doc, toc_items)
    add_executive_summary(doc)
    add_reading_guide(doc)
    add_document_control(doc)

    i = 0
    skip_until_content = True
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    table_buffer: list[list[str]] = []
    first_chapter = True
    converted_sections: list[str] = []
    stats = {"tables": 0, "code_blocks": 0, "headings_h2": 0, "headings_h3": 0}

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if skip_until_content:
            if stripped.startswith("## 1."):
                skip_until_content = False
            else:
                i += 1
                continue

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, code_lang)
                stats["code_blocks"] += 1
                code_lines = []
                code_lang = ""
                in_code = False
            else:
                in_code = True
                code_lang = stripped[3:].strip()
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            if is_table_separator(stripped):
                i += 1
                continue
            table_buffer.append(parse_table_row(stripped))
            if i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                i += 1
                continue
            add_table(doc, table_buffer)
            stats["tables"] += 1
            table_buffer = []
            i += 1
            continue

        if table_buffer:
            add_table(doc, table_buffer)
            stats["tables"] += 1
            table_buffer = []

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("# ") and not stripped.startswith("## "):
            i += 1
            continue

        if stripped.startswith("## "):
            title = stripped[3:].strip()
            if not first_chapter:
                doc.add_page_break()
            else:
                first_chapter = False
            doc.add_heading(title, level=1)
            converted_sections.append(title)
            if title in CHAPTER_NOTES:
                add_info_box(doc, "Uraian Singkat Bab", CHAPTER_NOTES[title])
            if title == "5. Kebutuhan Fungsional":
                add_info_box(
                    doc,
                    "Keterangan Tingkat Prioritas",
                    "Must — kebutuhan wajib pada rilis awal.  |  "
                    "Should — kebutuhan yang sangat direkomendasikan.  |  "
                    "Could — kebutuhan opsional yang dapat ditunda ke fase berikutnya.",
                    bg=COLOR_TIP_BG,
                )
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            stats["headings_h2"] += 1
            i += 1
            continue

        if stripped.startswith("#### "):
            title = stripped[5:].strip()
            if title.startswith("`") or re.match(r"^[a-z_]+$", title.replace("`", "")):
                p = doc.add_paragraph()
                run = p.add_run(f"Tabel: {title.strip('`')}")
                run.bold = True
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = COLOR_ACCENT
            else:
                doc.add_heading(title, level=3)
            stats["headings_h3"] += 1
            i += 1
            continue

        if stripped.startswith(">"):
            quote_text, i = collect_blockquote(lines, i)
            add_blockquote(doc, quote_text)
            continue

        if stripped.startswith("- "):
            add_bullet(doc, stripped[2:])
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            i += 1
            continue

        if is_persona_line(stripped):
            doc.add_heading(stripped.strip("*"), level=3)
            i += 1
            continue

        if is_label_line(stripped):
            p = doc.add_paragraph()
            parse_inline_formatting(p, stripped, base_size=11)
            p.runs[0].bold = True if p.runs else None
            i += 1
            continue

        p = doc.add_paragraph()
        parse_inline_formatting(p, stripped)
        i += 1

    if table_buffer:
        add_table(doc, table_buffer)
        stats["tables"] += 1

    doc.save(OUT_FILE)

    missing = [s for s in expected_sections if s not in converted_sections]
    print(f"Generated : {OUT_FILE}")
    print(f"Size      : {OUT_FILE.stat().st_size / 1024:.1f} KB")
    print(f"Sections  : {len(converted_sections)}/{len(expected_sections)} bab dikonversi")
    print(f"Tables    : {stats['tables']}")
    print(f"Code blocks: {stats['code_blocks']}")
    if missing:
        print(f"WARNING - Bab terlewat: {missing}")
    else:
        print("OK - Seluruh bab utama berhasil dikonversi.")


if __name__ == "__main__":
    convert_markdown_to_docx()
