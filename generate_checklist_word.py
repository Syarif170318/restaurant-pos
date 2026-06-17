# -*- coding: utf-8 -*-
"""Generate dokumen Word checklist manual testing dari Markdown."""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = Path(__file__).parent
MD_FILE = BASE_DIR / "RDP_CHECKLIST_MANUAL_TESTING.md"
OUT_FILE = BASE_DIR / "RDP_CHECKLIST_MANUAL_TESTING.docx"

COLOR_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
COLOR_SECONDARY = RGBColor(0x37, 0x41, 0x51)
COLOR_HEADER_BG = "1E40AF"
COLOR_ALT_ROW = "F1F5F9"
COLOR_ACCENT = RGBColor(0x0F, 0x76, 0x6E)


def set_cell_shading(cell, fill_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def add_header_footer(section):
    header = section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Manual Testing Checklist — Restaurant POS")
    hr.font.size = Pt(8)
    hr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    hr.font.italic = True

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Halaman ").font.size = Pt(9)
    run = fp.add_run()
    for el_type, text in [("begin", None), (None, "PAGE"), ("separate", None), ("end", None)]:
        if text:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), el_type)
            run._r.append(fld)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def setup_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_SECONDARY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color, space_before in [
        (1, 20, COLOR_PRIMARY, 0),
        (2, 14, COLOR_PRIMARY, 14),
        (3, 12, COLOR_ACCENT, 10),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.keep_with_next = True

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        add_header_footer(section)


def add_run_text(paragraph, text: str, bold=False, size=11, font="Calibri", color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|$", line.strip()))


def add_styled_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run_text(p, h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLOR_HEADER_BG)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            display = val if val else "☐ Ya  ☐ Tidak  ☐ N/A" if headers[c_idx] in ("✓", "Status") else ""
            if headers[c_idx] == "ID" and val.startswith("FR-"):
                add_run_text(p, display, bold=True, size=10, color=COLOR_PRIMARY)
            elif headers[c_idx] == "✓" or (headers[c_idx] == "Status" and not val):
                add_run_text(p, display, size=9)
            else:
                add_run_text(p, display, size=10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, COLOR_ALT_ROW)

    doc.add_paragraph()


def build_cover(doc: Document):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(title, "CHECKLIST", bold=True, size=26, color=COLOR_PRIMARY)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(subtitle, "Manual Testing — Restaurant POS System", bold=True, size=16, color=COLOR_SECONDARY)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_text(meta, "Berdasarkan RDP v1.1 + Phase 3 Full Features", size=11, color=COLOR_ACCENT)
    meta.paragraph_format.space_after = Pt(20)

    info = doc.add_table(rows=5, cols=2)
    info.style = "Table Grid"
    labels = [
        ("Proyek", "Restaurant POS (Dual-Client)"),
        ("URL Dev", "http://localhost:3001"),
        ("Mobile", "pos-mobile (Expo)"),
        ("Smoke Test", "33 skenario API (npm run smoke-test)"),
        ("Versi Dokumen", "Juni 2026"),
    ]
    for i, (label, value) in enumerate(labels):
        info.rows[i].cells[0].text = ""
        info.rows[i].cells[1].text = ""
        add_run_text(info.rows[i].cells[0].paragraphs[0], label, bold=True, size=10)
        add_run_text(info.rows[i].cells[1].paragraphs[0], value, size=10)
        set_cell_shading(info.rows[i].cells[0], COLOR_ALT_ROW)

    doc.add_paragraph()
    note = doc.add_paragraph()
    add_run_text(note, "Petunjuk: ", bold=True)
    add_run_text(
        note,
        "Centang kolom Hasil Uji dengan Ya / Tidak / N/A setelah setiap skenario diuji. "
        "Isi bagian penutup di akhir dokumen.",
    )

    doc.add_paragraph()
    acc = doc.add_heading("Akun Demo", level=2)
    _ = acc
    acc_table = doc.add_table(rows=6, cols=3)
    acc_table.style = "Table Grid"
    for i, h in enumerate(["User", "Password / PIN", "Peran"]):
        cell = acc_table.rows[0].cells[i]
        cell.text = ""
        add_run_text(cell.paragraphs[0], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLOR_HEADER_BG)
    accounts = [
        ("admin", "password123", "Administrator"),
        ("kasir1", "PIN 1111", "Kasir Outlet 1"),
        ("kasir2", "PIN 5555", "Kasir Outlet 2"),
        ("manager1", "PIN 3333", "Manager"),
        ("dapur1", "PIN 2222", "Kitchen"),
    ]
    for r, (u, p, role) in enumerate(accounts, start=1):
        acc_table.rows[r].cells[0].text = ""
        acc_table.rows[r].cells[1].text = ""
        acc_table.rows[r].cells[2].text = ""
        add_run_text(acc_table.rows[r].cells[0].paragraphs[0], u, size=10)
        add_run_text(acc_table.rows[r].cells[1].paragraphs[0], p, size=10)
        add_run_text(acc_table.rows[r].cells[2].paragraphs[0], role, size=10)

    doc.add_paragraph()
    steps = doc.add_heading("Cara Menjalankan Uji", level=2)
    _ = steps
    for step in [
        "cd pos-app && npm run dev          → Dashboard http://localhost:3001",
        "cd pos-mobile && npm start         → App kasir (emulator/HP)",
        "cd pos-app && npm run db:seed      → Reset data demo (opsional)",
        "cd pos-app && npm run smoke-test   → Cek API otomatis (33 skenario)",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        add_run_text(p, step, size=10, font="Consolas")

    doc.add_page_break()


def build_signoff(doc: Document):
    doc.add_heading("Penutup & Sign-Off", level=1)
    sign = doc.add_table(rows=6, cols=2)
    sign.style = "Table Grid"
    fields = [
        ("Nama Tester", "_______________________________"),
        ("Tanggal Uji", "_______________________________"),
        ("Lingkungan", "☐ Development   ☐ Staging   ☐ Production"),
        ("Total Passed", "_____ / _____ skenario"),
        ("Total Failed", "_____ skenario"),
        ("Disetujui oleh", "_______________________________"),
    ]
    for i, (label, value) in enumerate(fields):
        sign.rows[i].cells[0].text = ""
        sign.rows[i].cells[1].text = ""
        add_run_text(sign.rows[i].cells[0].paragraphs[0], label, bold=True, size=11)
        add_run_text(sign.rows[i].cells[1].paragraphs[0], value, size=11)
        set_cell_shading(sign.rows[i].cells[0], COLOR_ALT_ROW)

    doc.add_paragraph()
    doc.add_heading("Catatan Temuan / Bug", level=2)
    bug_table = doc.add_table(rows=6, cols=4)
    bug_table.style = "Table Grid"
    for i, h in enumerate(["No", "FR-ID / Fitur", "Deskripsi Masalah", "Severity"]):
        cell = bug_table.rows[0].cells[i]
        cell.text = ""
        add_run_text(cell.paragraphs[0], h, bold=True, size=10, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, COLOR_HEADER_BG)
    for r in range(1, 6):
        for c in range(4):
            bug_table.rows[r].cells[c].text = ""
            add_run_text(bug_table.rows[r].cells[c].paragraphs[0], str(r) if c == 0 else "", size=10)


def parse_markdown(md_text: str) -> Document:
    doc = Document()
    setup_styles(doc)
    build_cover(doc)

    lines = md_text.splitlines()
    i = 0
    headers: list[str] = []
    rows: list[list[str]] = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        if line.startswith("# ") and not line.startswith("## "):
            i += 1
            continue

        if line.startswith("## "):
            if in_table and headers and rows:
                add_styled_table(doc, headers, rows)
                headers, rows, in_table = [], [], False
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.strip().startswith("|") and not is_table_separator(line):
            cells = parse_table_row(line)
            if not in_table:
                headers = cells
                in_table = True
            else:
                rows.append(cells)
            i += 1
            continue

        if is_table_separator(line):
            i += 1
            continue

        if line.strip() == "---":
            if in_table and headers and rows:
                add_styled_table(doc, headers, rows)
                headers, rows, in_table = [], [], False
            i += 1
            continue

        if line.strip().startswith("**Tester:**") or line.strip().startswith("**Tanggal:**"):
            i += 1
            continue

        stripped = line.strip()
        if stripped and not stripped.startswith("```"):
            if in_table and headers and rows:
                add_styled_table(doc, headers, rows)
                headers, rows, in_table = [], [], False
            if stripped.startswith("**") and stripped.endswith("**"):
                p = doc.add_paragraph()
                add_run_text(p, stripped.strip("*"), bold=True)
            elif not stripped.startswith("cd ") and "npm run" not in stripped:
                p = doc.add_paragraph()
                add_run_text(p, stripped)

        if stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph()
                add_run_text(p, lines[i], size=9, font="Consolas")
                i += 1

        i += 1

    if in_table and headers and rows:
        add_styled_table(doc, headers, rows)

    build_signoff(doc)
    return doc


def main():
    if not MD_FILE.exists():
        raise FileNotFoundError(f"Markdown tidak ditemukan: {MD_FILE}")

    md_text = MD_FILE.read_text(encoding="utf-8")
    doc = parse_markdown(md_text)
    doc.save(OUT_FILE)

    size_kb = OUT_FILE.stat().st_size / 1024
    table_count = len(doc.tables)
    print(f"Generated : {OUT_FILE}")
    print(f"Size      : {size_kb:.1f} KB")
    print(f"Tables    : {table_count}")
    print("OK - Checklist manual testing berhasil dikonversi ke Word.")


if __name__ == "__main__":
    main()
